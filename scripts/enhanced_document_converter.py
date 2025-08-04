#!/usr/bin/env python3
"""
Enhanced Document to Markdown Converter
=======================================

Advanced converter for research documents with comprehensive formatting support:
- PDF files (academic papers, handbooks) with table extraction
- DOCX files (research proposals, manuscripts) with full structure preservation
- Tables, lists, headers, formatting preservation
- Batch processing for entire directories
- Intelligent text cleaning and markdown formatting

Author: Alex Finch (AI Research Partner)
Project: AIRS - AI Readiness Score Development
Purpose: Enable systematic analysis of 100+ research documents with formatting integrity
Version: 2.0 - Enhanced Formatting Support
"""

import os
import sys
import argparse
from pathlib import Path
import re
from typing import Optional, List, Dict, Tuple, Any
import logging

# PDF processing imports
try:
    import pdfplumber
    import PyPDF2
    import fitz  # PyMuPDF for advanced table extraction
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("⚠️  PDF libraries not available. Install with: pip install pdfplumber PyPDF2 pymupdf")

# DOCX processing imports
try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
    import docx2txt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  DOCX libraries not available. Install with: pip install python-docx docx2txt")

# Enhanced formatting imports
try:
    from tabulate import tabulate
    import html2text
    FORMATTING_AVAILABLE = True
except ImportError:
    FORMATTING_AVAILABLE = False
    print("⚠️  Enhanced formatting libraries not available. Install with: pip install tabulate html2text")

class EnhancedDocumentConverter:
    """
    Advanced multi-format document converter with comprehensive formatting preservation.
    
    Enhanced Features:
    - Table extraction and markdown formatting for both PDF and DOCX
    - List preservation (numbered, bulleted, nested)
    - Header hierarchy maintenance
    - Bold, italic, and other text formatting
    - Image placeholder handling
    - Citation and reference formatting
    - Error logging and detailed reporting
    """
    
    def __init__(self, output_dir: str = None, preserve_formatting: bool = True):
        self.output_dir = Path(output_dir) if output_dir else None
        if self.output_dir:
            self.output_dir.mkdir(exist_ok=True)
        self.preserve_formatting = preserve_formatting
        self.conversion_log = []
        
        # Configure logging
        logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
        self.logger = logging.getLogger(__name__)
        
    def convert_pdf_enhanced(self, pdf_path: str) -> Tuple[bool, str, str]:
        """
        Enhanced PDF conversion with high-fidelity table and formatting extraction.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Tuple of (success, markdown_content, error_message)
        """
        if not PDF_AVAILABLE:
            return False, "", "PDF libraries not installed"
            
        pdf_path = Path(pdf_path)
        markdown_content = f"# {pdf_path.stem}\n\n"
        
        try:
            # Primary method: pdfplumber for high-fidelity extraction
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    markdown_content += f"\n---\n## Page {page_num}\n\n"
                    
                    # Strategy 1: Extract tables with better settings
                    table_settings = {
                        "vertical_strategy": "lines_strict",
                        "horizontal_strategy": "lines_strict",
                        "snap_tolerance": 3,
                        "join_tolerance": 3,
                        "edge_min_length": 3,
                        "min_words_vertical": 3,
                        "min_words_horizontal": 1,
                        "text_tolerance": 3,
                        "text_x_tolerance": 3,
                        "text_y_tolerance": 3,
                    }
                    
                    tables = page.extract_tables(table_settings)
                    processed_table_areas = []
                    
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            if table and len(table) > 0 and self._is_valid_table(table):
                                # Get table bounding box to exclude from text extraction
                                table_bbox = page.within_bbox(page.bbox).extract_tables()[table_num-1] if table_num <= len(page.extract_tables()) else None
                                
                                markdown_content += f"\n### Table {table_num}\n\n"
                                formatted_table = self._format_table_to_markdown_enhanced(table)
                                if formatted_table:
                                    markdown_content += formatted_table + "\n\n"
                                else:
                                    # Fallback to simple formatting
                                    markdown_content += self._format_table_simple_list(table) + "\n\n"
                    
                    # Strategy 2: Extract text outside of table areas
                    text = page.extract_text()
                    if text:
                        # Clean and format text with better structure detection
                        formatted_text = self._enhance_text_formatting_advanced(text)
                        markdown_content += formatted_text + "\n\n"
                        
            if len(markdown_content) > len(f"# {pdf_path.stem}\n\n"):
                return True, self._clean_text_enhanced(markdown_content), ""
                
        except Exception as e:
            self.logger.warning(f"pdfplumber enhanced failed for {pdf_path.name}: {e}")
            
        # Fallback to PyMuPDF with better table detection
        try:
            return self._pymupdf_enhanced_extraction(pdf_path)
        except Exception as e:
            self.logger.warning(f"PyMuPDF enhanced failed for {pdf_path.name}: {e}")
            
        # Final fallback: Basic PyPDF2
        return self._basic_pdf_conversion(pdf_path)
    
    def convert_docx_enhanced(self, docx_path: str) -> Tuple[bool, str, str]:
        """
        Enhanced DOCX conversion with high-fidelity formatting preservation.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Tuple of (success, markdown_content, error_message)
        """
        if not DOCX_AVAILABLE:
            return False, "", "DOCX libraries not installed"
            
        docx_path = Path(docx_path)
        
        try:
            # Primary method: python-docx with comprehensive structure analysis
            doc = Document(docx_path)
            markdown_content = f"# {docx_path.stem}\n\n"
            
            # Process document elements in strict order for structure preservation
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = Paragraph(element, doc)
                    md_text = self._process_paragraph_enhanced(paragraph)
                    if md_text.strip():
                        markdown_content += md_text + "\n\n"
                        
                elif element.tag.endswith('tbl'):  # Table
                    # Find the corresponding table object and process with high fidelity
                    for table in doc.tables:
                        if table._element == element:
                            md_table = self._process_docx_table_enhanced(table)
                            if md_table:
                                markdown_content += md_table + "\n\n"
                            break
            
            # Process any remaining tables that weren't caught in the element loop
            for table in doc.tables:
                if table._element not in [elem for elem in doc.element.body if elem.tag.endswith('tbl')]:
                    md_table = self._process_docx_table_enhanced(table)
                    if md_table:
                        markdown_content += "\n### Additional Table\n\n" + md_table + "\n\n"
                            
            return True, self._clean_text_enhanced(markdown_content), ""
            
        except Exception as e:
            self.logger.warning(f"python-docx enhanced failed for {docx_path.name}: {e}")
            
        # Fallback to basic conversion with improvements
        return self._basic_docx_conversion_enhanced(docx_path)
    
    def _format_table_to_markdown_enhanced(self, table: List[List[str]]) -> str:
        """Enhanced table formatting with better structure detection and cleaning."""
        if not table or len(table) < 1:
            return ""
            
        try:
            # Advanced table cleaning and normalization
            cleaned_table = []
            max_cols = 0
            
            # First pass: identify actual table structure
            for row_idx, row in enumerate(table):
                if row:
                    # Clean each cell more aggressively
                    cleaned_row = []
                    for cell in row:
                        if cell is not None:
                            # Enhanced cell cleaning
                            cell_text = str(cell).strip()
                            cell_text = re.sub(r'\s+', ' ', cell_text)  # Normalize whitespace
                            cell_text = re.sub(r'[\r\n]+', ' ', cell_text)  # Remove line breaks
                            cell_text = cell_text.replace('|', '\\|')  # Escape pipe characters
                            
                            # Keep cell if it has content or if we're building a structured row
                            if cell_text or (len(cleaned_row) > 0 and any(cleaned_row)):
                                cleaned_row.append(cell_text)
                    
                    # Only keep rows that have meaningful content
                    if cleaned_row and any(cell.strip() for cell in cleaned_row):
                        # Detect if this looks like a header row (be more conservative)
                        is_header = (row_idx == 0 and
                                   any(keyword in ' '.join(cleaned_row).lower() 
                                       for keyword in ['course', 'number', 'title', 'credit', 'hours', 'name', 'description', 'column', 'field', 'item', 'type']))
                        
                        cleaned_table.append((cleaned_row, is_header))
                        max_cols = max(max_cols, len(cleaned_row))
            
            if not cleaned_table or max_cols < 2:
                return self._format_table_simple_list(table)
            
            # Normalize column count and identify header rows
            normalized_table = []
            header_indices = []
            
            for idx, (row, is_header) in enumerate(cleaned_table):
                # Pad row to max columns
                while len(row) < max_cols:
                    row.append("")
                normalized_table.append(row)
                if is_header:
                    header_indices.append(idx)
            
            # Generate markdown table
            if FORMATTING_AVAILABLE and len(normalized_table) > 1:
                # Use first row as header if no explicit headers found
                if not header_indices:
                    header_indices = [0]
                
                header_row = normalized_table[header_indices[0]]
                data_rows = [row for idx, row in enumerate(normalized_table) if idx not in header_indices]
                
                if data_rows:
                    return tabulate(data_rows, headers=header_row, tablefmt="pipe", numalign="left", stralign="left")
                else:
                    # Only header row
                    return self._create_simple_markdown_table([header_row])
            else:
                return self._create_simple_markdown_table(normalized_table)
                
        except Exception as e:
            self.logger.warning(f"Enhanced table formatting failed: {e}")
            return self._format_table_simple_list(table)
    
    def _create_simple_markdown_table(self, table: List[List[str]]) -> str:
        """Create a simple, clean markdown table."""
        if not table:
            return ""
        
        markdown_lines = []
        
        # Header row
        header = "| " + " | ".join(cell.strip() for cell in table[0]) + " |"
        markdown_lines.append(header)
        
        # Separator row
        separator = "| " + " | ".join("---" for _ in table[0]) + " |"
        markdown_lines.append(separator)
        
        # Data rows
        for row in table[1:]:
            row_text = "| " + " | ".join(cell.strip() for cell in row) + " |"
            markdown_lines.append(row_text)
        
        return "\n".join(markdown_lines)
    
    def _format_table_simple_list(self, table: List[List[str]]) -> str:
        """Format complex tables as structured lists when table format fails."""
        if not table:
            return ""
        
        content = ""
        for i, row in enumerate(table):
            if row and any(str(cell).strip() for cell in row):
                non_empty_cells = [str(cell).strip() for cell in row if str(cell).strip()]
                if i == 0:
                    content += "**" + " | ".join(non_empty_cells) + "**\n\n"
                else:
                    content += "- " + " | ".join(non_empty_cells) + "\n"
        
        return content
    
    def _manual_table_format(self, table: List[List[str]]) -> str:
        """Improved manual markdown table formatting."""
        if not table or len(table) < 1:
            return ""
            
        markdown_table = ""
        
        try:
            # Ensure consistent column count
            max_cols = max(len(row) for row in table)
            
            # Header row
            if len(table) > 0:
                header_row = table[0]
                while len(header_row) < max_cols:
                    header_row.append("")
                
                header = "| " + " | ".join(cell.strip() for cell in header_row) + " |\n"
                separator = "| " + " | ".join("---" for _ in range(max_cols)) + " |\n"
                markdown_table += header + separator
                
                # Data rows
                for row in table[1:]:
                    while len(row) < max_cols:
                        row.append("")
                    row_text = "| " + " | ".join(cell.strip() for cell in row) + " |\n"
                    markdown_table += row_text
                    
        except Exception as e:
            self.logger.warning(f"Manual table formatting failed: {e}")
            return self._manual_table_format_simple(table)
                    
        return markdown_table
    
    def _manual_table_format_simple(self, table: List[List[str]]) -> str:
        """Simple fallback table formatting."""
        if not table:
            return ""
        
        # Just format as simple list if table is too complex
        content = ""
        for i, row in enumerate(table):
            if row and any(str(cell).strip() for cell in row):
                if i == 0:
                    content += "**Headers:** " + " | ".join(str(cell).strip() for cell in row if str(cell).strip()) + "\n\n"
                else:
                    content += "- " + " | ".join(str(cell).strip() for cell in row if str(cell).strip()) + "\n"
        
        return content
    
    def _process_paragraph(self, paragraph: Paragraph) -> str:
        """Process DOCX paragraph with formatting preservation."""
        if not paragraph.text.strip():
            return ""
            
        # Determine paragraph style
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        text = paragraph.text.strip()
        
        # Handle different paragraph types
        if 'heading 1' in style_name or 'title' in style_name:
            return f"# {text}"
        elif 'heading 2' in style_name:
            return f"## {text}"
        elif 'heading 3' in style_name:
            return f"### {text}"
        elif 'heading 4' in style_name:
            return f"#### {text}"
        elif 'heading 5' in style_name:
            return f"##### {text}"
        elif 'heading 6' in style_name:
            return f"###### {text}"
        elif 'list' in style_name:
            return f"- {text}"
        elif 'quote' in style_name or 'citation' in style_name:
            return f"> {text}"
        else:
            # Process inline formatting
            return self._process_inline_formatting(paragraph)
    
    def _process_paragraph_enhanced(self, paragraph: Paragraph) -> str:
        """Enhanced paragraph processing with better formatting detection."""
        if not paragraph.text.strip():
            return ""
            
        # Get style information
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        text = paragraph.text.strip()
        
        # Enhanced style detection
        if any(heading in style_name for heading in ['title', 'heading 1']):
            return f"# {text}"
        elif 'heading 2' in style_name:
            return f"## {text}"
        elif 'heading 3' in style_name:
            return f"### {text}"
        elif 'heading 4' in style_name:
            return f"#### {text}"
        elif 'heading 5' in style_name:
            return f"##### {text}"
        elif 'heading 6' in style_name:
            return f"###### {text}"
        elif any(list_style in style_name for list_style in ['list', 'bullet']):
            return f"- {text}"
        elif any(quote_style in style_name for quote_style in ['quote', 'citation', 'block']):
            return f"> {text}"
        else:
            # Process with enhanced inline formatting
            return self._process_inline_formatting_enhanced(paragraph)
    
    def _process_inline_formatting(self, paragraph: Paragraph) -> str:
        """Extract inline formatting (bold, italic) from paragraph."""
        text = ""
        
        for run in paragraph.runs:
            run_text = run.text
            
            # Apply formatting
            if run.bold and run.italic:
                run_text = f"***{run_text}***"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"*{run_text}*"
                
            text += run_text
            
        return text
    
    def _process_inline_formatting_enhanced(self, paragraph: Paragraph) -> str:
        """Enhanced inline formatting extraction with better text processing."""
        text = ""
        
        # Handle paragraphs with no runs (text stored directly in paragraph)
        if not paragraph.runs:
            return paragraph.text
        
        for run in paragraph.runs:
            run_text = run.text
            
            # Apply comprehensive formatting
            formatting = []
            if getattr(run, 'bold', False):
                formatting.append('**')
            if getattr(run, 'italic', False):
                formatting.append('*')
            if getattr(run.font, 'superscript', False):
                run_text = f"^{run_text}^"
            if getattr(run.font, 'subscript', False):
                run_text = f"_{run_text}_"
            
            # Apply formatting
            if formatting:
                format_chars = ''.join(formatting)
                run_text = f"{format_chars}{run_text}{format_chars[::-1]}"
                
            text += run_text
            
        return text
    
    def _process_docx_table(self, table: Table) -> str:
        """Process DOCX table and convert to markdown."""
        try:
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get cell text and clean it
                    cell_text = cell.text.strip().replace('\n', ' ').replace('\r', '')
                    row_data.append(cell_text)
                table_data.append(row_data)
                
            return self._format_table_to_markdown(table_data)
            
        except Exception as e:
            self.logger.warning(f"DOCX table processing failed: {e}")
            return ""
    
    def _process_docx_table_enhanced(self, table: Table) -> str:
        """Enhanced DOCX table processing with better structure preservation."""
        try:
            table_data = []
            
            # Extract table with better cell handling
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    # Enhanced cell text extraction
                    cell_text = ""
                    for paragraph in cell.paragraphs:
                        para_text = self._process_inline_formatting_enhanced(paragraph)
                        if para_text.strip():
                            cell_text += para_text.strip() + " "
                    
                    # Clean cell text
                    cell_text = cell_text.strip()
                    cell_text = re.sub(r'\s+', ' ', cell_text)
                    cell_text = cell_text.replace('|', '\\|')  # Escape pipes
                    
                    row_data.append(cell_text)
                
                if any(cell.strip() for cell in row_data):  # Only add non-empty rows
                    table_data.append(row_data)
                
            return self._format_table_to_markdown_enhanced(table_data)
            
        except Exception as e:
            self.logger.warning(f"Enhanced DOCX table processing failed: {e}")
            return ""
    
    def _process_pdf_blocks(self, blocks: Dict) -> str:
        """Process PDF text blocks with formatting information."""
        content = ""
        
        try:
            for block in blocks.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            flags = span.get("flags", 0)
                            
                            # Apply formatting based on flags
                            if flags & 2**4:  # Bold
                                text = f"**{text}**"
                            if flags & 2**1:  # Italic
                                text = f"*{text}*"
                                
                            line_text += text
                            
                        if line_text.strip():
                            content += line_text + "\n"
                            
        except Exception as e:
            self.logger.warning(f"PDF block processing failed: {e}")
            
        return content
    
    def _process_pdf_blocks_enhanced(self, blocks: Dict) -> str:
        """Enhanced PDF block processing with better formatting detection."""
        content = ""
        
        try:
            for block in blocks.get("blocks", []):
                if block.get("type") == 0:  # Text block
                    block_text = ""
                    for line in block.get("lines", []):
                        line_text = ""
                        line_formats = []
                        
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            flags = span.get("flags", 0)
                            font_size = span.get("size", 12)
                            
                            # Enhanced formatting detection
                            if flags & 16:  # Bold
                                text = f"**{text}**"
                            if flags & 2:  # Italic
                                text = f"*{text}*"
                            if font_size > 14:  # Likely header
                                line_formats.append("header")
                                
                            line_text += text
                        
                        if line_text.strip():
                            # Apply line-level formatting
                            if "header" in line_formats:
                                line_text = f"## {line_text.strip()}"
                            
                            block_text += line_text + "\n"
                    
                    if block_text.strip():
                        content += block_text + "\n"
                            
        except Exception as e:
            self.logger.warning(f"Enhanced PDF block processing failed: {e}")
            
        return content
    
    def _enhance_text_formatting(self, text: str) -> str:
        """Enhance plain text with markdown formatting detection."""
        if not text:
            return ""
            
        lines = text.split('\n')
        enhanced_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                enhanced_lines.append("")
                continue
                
            # Detect potential headers (all caps, short lines)
            if (len(line) < 100 and 
                line.isupper() and 
                not any(char.isdigit() for char in line[:10])):
                enhanced_lines.append(f"## {line.title()}")
            
            # Detect numbered lists
            elif re.match(r'^\d+\.?\s+', line):
                cleaned_line = re.sub(r'^\d+\.?\s+', '', line)
                enhanced_lines.append(f"1. {cleaned_line}")
            
            # Detect bullet points
            elif re.match(r'^[•·▪▫-]\s+', line):
                cleaned_line = re.sub(r'^[•·▪▫-]\s+', '', line)
                enhanced_lines.append(f"- {cleaned_line}")
            
            else:
                enhanced_lines.append(line)
                
        return '\n'.join(enhanced_lines)
    
    def _enhance_text_formatting_advanced(self, text: str) -> str:
        """Advanced text formatting with better structure detection."""
        if not text:
            return ""
            
        lines = text.split('\n')
        enhanced_lines = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                enhanced_lines.append("")
                continue
            
            # Enhanced header detection
            if self._is_likely_header(line, i, lines):
                level = self._determine_header_level(line)
                enhanced_lines.append(f"{'#' * level} {line}")
            
            # Enhanced list detection
            elif re.match(r'^\d+[\.)]\s+', line):
                cleaned_line = re.sub(r'^\d+[\.)]\s+', '', line)
                enhanced_lines.append(f"1. {cleaned_line}")
            
            elif re.match(r'^[•·▪▫○●-]\s+', line):
                cleaned_line = re.sub(r'^[•·▪▫○●-]\s+', '', line)
                enhanced_lines.append(f"- {cleaned_line}")
            
            # Detect emphasis patterns
            elif line.isupper() and len(line) < 80 and len(line) > 3:
                enhanced_lines.append(f"**{line.title()}**")
            
            else:
                enhanced_lines.append(line)
                
        return '\n'.join(enhanced_lines)
    
    def _is_likely_header(self, line: str, index: int, all_lines: List[str]) -> bool:
        """Determine if a line is likely a header."""
        # Check various header indicators
        if len(line) > 100:  # Too long for header
            return False
        
        # Check if followed by content
        has_following_content = (index + 1 < len(all_lines) and 
                               all_lines[index + 1].strip() and
                               not all_lines[index + 1].strip().isupper())
        
        # Header patterns
        is_all_caps = line.isupper()
        is_title_case = line.istitle()
        has_header_keywords = any(word in line.lower() for word in 
                                ['chapter', 'section', 'part', 'overview', 'introduction', 'conclusion'])
        is_short = len(line) < 80
        
        return (is_all_caps or is_title_case or has_header_keywords) and is_short and has_following_content
    
    def _determine_header_level(self, line: str) -> int:
        """Determine appropriate header level based on content."""
        if any(word in line.lower() for word in ['chapter', 'part']):
            return 1
        elif any(word in line.lower() for word in ['section', 'overview']):
            return 2
        elif line.isupper():
            return 2
        elif line.istitle():
            return 3
        else:
            return 4
    
    def _pymupdf_enhanced_extraction(self, pdf_path: Path) -> Tuple[bool, str, str]:
        """Enhanced PyMuPDF extraction with better table detection."""
        try:
            markdown_content = f"# {pdf_path.stem}\n\n"
            doc = fitz.open(pdf_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                markdown_content += f"\n---\n## Page {page_num + 1}\n\n"
                
                # Try to extract tables using PyMuPDF
                try:
                    tables = page.find_tables()
                    for table_num, table in enumerate(tables, 1):
                        table_data = table.extract()
                        if table_data and self._is_valid_table(table_data):
                            markdown_content += f"\n### Table {table_num}\n\n"
                            formatted_table = self._format_table_to_markdown_enhanced(table_data)
                            if formatted_table:
                                markdown_content += formatted_table + "\n\n"
                except:
                    # Fallback to text blocks
                    pass
                
                # Extract text blocks with formatting
                blocks = page.get_text("dict")
                formatted_blocks = self._process_pdf_blocks_enhanced(blocks)
                markdown_content += formatted_blocks + "\n\n"
                
            doc.close()
            
            if len(markdown_content) > len(f"# {pdf_path.stem}\n\n"):
                return True, self._clean_text_enhanced(markdown_content), ""
            else:
                return False, "", "No content extracted"
                
        except Exception as e:
            return False, "", f"PyMuPDF enhanced extraction failed: {e}"
    
    def _basic_pdf_conversion(self, pdf_path: Path) -> Tuple[bool, str, str]:
        """Basic PDF conversion fallback."""
        try:
            markdown_content = f"# {pdf_path.stem}\n\n"
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text:
                        markdown_content += f"\n## Page {page_num}\n\n{text}\n\n"
                        
            return True, self._clean_text_enhanced(markdown_content), ""
            
        except Exception as e:
            return False, "", f"All PDF extraction methods failed: {e}"
    
    def _basic_docx_conversion(self, docx_path: Path) -> Tuple[bool, str, str]:
        """Basic DOCX conversion fallback."""
        try:
            text = docx2txt.process(docx_path)
            if text:
                markdown_content = f"# {docx_path.stem}\n\n{text}"
                return True, self._clean_text_enhanced(markdown_content), ""
                
        except Exception as e:
            return False, "", f"All DOCX extraction methods failed: {e}"
            
        return False, "", "No text content extracted"
    
    def _basic_docx_conversion_enhanced(self, docx_path: Path) -> Tuple[bool, str, str]:
        """Enhanced basic DOCX conversion with better structure preservation."""
        try:
            # Use docx2txt but with post-processing
            text = docx2txt.process(docx_path)
            if text:
                markdown_content = f"# {docx_path.stem}\n\n"
                
                # Apply advanced text formatting
                formatted_text = self._enhance_text_formatting_advanced(text)
                markdown_content += formatted_text
                
                return True, self._clean_text_enhanced(markdown_content), ""
                
        except Exception as e:
            return False, "", f"Enhanced DOCX conversion failed: {e}"
            
        return False, "", "No text content extracted"
    
    def _clean_text_enhanced(self, text: str) -> str:
        """Enhanced text cleaning with formatting preservation and duplicate removal."""
        # Remove excessive whitespace but preserve structure
        text = re.sub(r'\n{4,}', '\n\n\n', text)  # Max 3 consecutive newlines
        text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces to single
        
        # Fix common extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # camelCase spacing
        text = re.sub(r'(\.)([A-Z])', r'\1 \2', text)     # Period spacing
        
        # Preserve academic citations
        text = re.sub(r'\(\s*([A-Za-z]+)\s*,\s*(\d{4})\s*\)', r'(\1, \2)', text)
        
        # Clean markdown formatting conflicts
        text = re.sub(r'\*{4,}', '***', text)  # Fix excessive asterisks
        text = re.sub(r'#{7,}', '######', text)  # Max 6 header levels
        
        # Fix malformed table formatting - more aggressive cleaning
        text = re.sub(r'\|\s*\|\s*\|', '| |', text)  # Empty table cells
        text = re.sub(r'\|(\s*\|){3,}', '| |', text)  # Multiple empty columns
        
        # Enhanced duplicate content removal
        lines = text.split('\n')
        cleaned_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip malformed table rows (too many pipes or separators)
            if line.count('|') > 10 and line.count(':---') > 5:
                i += 1
                continue
            
            # Detect and clean up table sections
            if '|' in line and ('---' in line or any('---' in lines[j] for j in range(max(0, i-2), min(len(lines), i+3)))):
                # We're in a table section - clean it up
                table_section = []
                table_start = i
                
                # Collect the table section
                while i < len(lines) and ('|' in lines[i] or not lines[i].strip()):
                    if lines[i].strip():
                        table_section.append(lines[i])
                    i += 1
                
                # Clean the table section
                if table_section:
                    cleaned_table = self._clean_table_section(table_section)
                    cleaned_lines.extend(cleaned_table)
                
                # Skip duplicate plain text that follows
                duplicate_skip_count = 0
                while (i < len(lines) and 
                       duplicate_skip_count < 15 and  # Look ahead max 15 lines
                       not lines[i].startswith('#') and
                       not lines[i].startswith('---')):
                    
                    line_content = lines[i].strip()
                    if (line_content and 
                        any(keyword in line_content.lower() 
                            for keyword in ['dba', 'course', 'credit', 'number', 'title', 'semester']) and
                        not line_content.startswith('*') and
                        not line_content.startswith('-') and
                        '|' not in line_content):
                        # This looks like duplicate table content
                        i += 1
                        duplicate_skip_count += 1
                    else:
                        break
                
                continue
            
            cleaned_lines.append(lines[i])
            i += 1
        
        # Remove consecutive duplicate lines
        final_lines = []
        prev_line = ""
        
        for line in cleaned_lines:
            if line.strip() != prev_line.strip() or not line.strip():
                final_lines.append(line)
            prev_line = line
        
        text = '\n'.join(final_lines)
        
        return text.strip()
    
    def _clean_table_section(self, table_lines: List[str]) -> List[str]:
        """Clean a table section to remove malformed rows."""
        cleaned = []
        
        for line in table_lines:
            # Skip obviously malformed table rows
            if line.count('|') > 15:  # Too many columns
                continue
            if line.count(':---') > 8:  # Too many separators
                continue
            if re.match(r'^\|\s*\|\s*\|\s*\|\s*\|', line):  # Too many empty columns at start
                continue
            
            # Clean up the line
            line = re.sub(r'\|\s*\|\s*\|\s*\|', '| |', line)  # Collapse empty columns
            line = re.sub(r'\|{2,}', '|', line)  # Remove duplicate pipes
            
            if line.strip():
                cleaned.append(line)
        
        return cleaned
    
    def convert_document(self, file_path: str) -> bool:
        """Convert single document with enhanced formatting support."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"❌ File not found: {file_path}")
            return False
            
        suffix = file_path.suffix.lower()
        
        print(f"🔄 Converting {file_path.name} with enhanced formatting...")
        
        if suffix == '.pdf':
            success, content, error = self.convert_pdf_enhanced(file_path)
        elif suffix in ['.docx', '.doc']:
            success, content, error = self.convert_docx_enhanced(file_path)
        else:
            print(f"❌ Unsupported file format: {suffix}")
            return False
            
        if not success:
            print(f"❌ Enhanced conversion failed for {file_path.name}: {error}")
            self.conversion_log.append({
                'file': str(file_path),
                'status': 'failed',
                'error': error,
                'enhanced': True
            })
            return False
            
        # Save with enhanced formatting notation
        if self.output_dir:
            output_file = self.output_dir / f"{file_path.stem}.md"
        else:
            output_file = file_path.parent / f"{file_path.stem}.md"
            
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(content)
                
            print(f"✅ Enhanced conversion: {file_path.name} → {output_file.name}")
            self.conversion_log.append({
                'file': str(file_path),
                'output': str(output_file),
                'status': 'success',
                'size': len(content),
                'enhanced': True,
                'tables_detected': content.count('|'),
                'headers_detected': content.count('#')
            })
            return True
            
        except Exception as e:
            print(f"❌ Failed to save {output_file}: {e}")
            return False
    
    def batch_convert(self, directory: str, file_pattern: str = "*") -> Dict[str, int]:
        """Batch convert with enhanced formatting."""
        directory = Path(directory)
        
        if not directory.exists():
            print(f"❌ Directory not found: {directory}")
            return {'total': 0, 'success': 0, 'failed': 0}
            
        supported_extensions = ['.pdf', '.docx', '.doc']
        files = []
        
        for ext in supported_extensions:
            files.extend(directory.rglob(f"{file_pattern}{ext}"))
            
        print(f"\n🔍 Found {len(files)} documents for enhanced conversion in {directory}")
        
        stats = {'total': len(files), 'success': 0, 'failed': 0}
        
        for file_path in files:
            if self.convert_document(file_path):
                stats['success'] += 1
            else:
                stats['failed'] += 1
                
        return stats
    
    def print_enhanced_summary(self):
        """Print enhanced conversion summary with formatting statistics."""
        if not self.conversion_log:
            print("\n📊 No conversions performed.")
            return
            
        successful = [entry for entry in self.conversion_log if entry['status'] == 'success']
        failed = [entry for entry in self.conversion_log if entry['status'] == 'failed']
        
        print(f"\n📊 Enhanced Conversion Summary:")
        print(f"   ✅ Successful: {len(successful)}")
        print(f"   ❌ Failed: {len(failed)}")
        
        if successful:
            total_size = sum(entry['size'] for entry in successful)
            total_tables = sum(entry.get('tables_detected', 0) for entry in successful)
            total_headers = sum(entry.get('headers_detected', 0) for entry in successful)
            
            print(f"   📝 Total content: {total_size:,} characters")
            print(f"   📊 Tables processed: {total_tables}")
            print(f"   📑 Headers detected: {total_headers}")
            
            if self.output_dir:
                print(f"   📁 Output directory: {self.output_dir}")
            else:
                print(f"   📁 Files saved alongside originals")
                
        if failed:
            print(f"\n❌ Failed conversions:")
            for entry in failed[:5]:
                print(f"   - {Path(entry['file']).name}: {entry['error']}")
            if len(failed) > 5:
                print(f"   ... and {len(failed) - 5} more")
    
    def _is_valid_table(self, table: List[List[str]]) -> bool:
        """Validate if extracted table is actually a meaningful table."""
        if not table or len(table) < 2:  # Need at least header + 1 data row
            return False
        
        # Check if table has reasonable structure
        row_lengths = [len(row) for row in table if row]
        if not row_lengths or max(row_lengths) < 2:  # Need at least 2 columns
            return False
        
        # Check for content density - avoid tables that are mostly empty
        total_cells = sum(row_lengths)
        filled_cells = 0
        
        for row in table:
            for cell in row:
                if cell and str(cell).strip():
                    filled_cells += 1
        
        # Require at least 30% filled cells
        if total_cells > 0 and (filled_cells / total_cells) < 0.3:
            return False
        
        # Check if it's likely a formatting artifact (all single characters, etc.)
        content_chars = 0
        for row in table:
            for cell in row:
                if cell and str(cell).strip():
                    content_chars += len(str(cell).strip())
        
        # Require reasonable content length
        if content_chars < 10:
            return False
        
        return True

def main():
    """Enhanced command-line interface."""
    parser = argparse.ArgumentParser(
        description="Enhanced Document Converter with comprehensive formatting support",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Enhanced Features:
  ✅ Table extraction and markdown formatting
  ✅ Header hierarchy preservation  
  ✅ Bold, italic, and text formatting
  ✅ List structure maintenance
  ✅ Citation and reference handling

Examples:
  # Enhanced single file conversion
  python enhanced_document_converter.py file.pdf
  
  # Enhanced batch conversion with formatting
  python enhanced_document_converter.py --batch papers/
  
  # Convert with formatting statistics
  python enhanced_document_converter.py --batch papers/ --pattern "*AIRS*"
        """
    )
    
    parser.add_argument('input', nargs='?', help='Input file or directory')
    parser.add_argument('--batch', metavar='DIR', help='Batch convert directory')
    parser.add_argument('--output', metavar='DIR', default=None, 
                       help='Output directory (default: same as source file)')
    parser.add_argument('--pattern', metavar='PATTERN', default='*',
                       help='File pattern for batch mode (default: *)')
    parser.add_argument('--basic', action='store_true',
                       help='Use basic conversion (faster, less formatting)')
    
    args = parser.parse_args()
    
    if not args.input and not args.batch:
        parser.print_help()
        return
        
    # Check dependencies
    if not PDF_AVAILABLE and not DOCX_AVAILABLE:
        print("❌ No document processing libraries available!")
        print("📦 Install with: pip install -r requirements.txt")
        return
        
    # Initialize converter
    preserve_formatting = not args.basic
    converter = EnhancedDocumentConverter(args.output, preserve_formatting)
    
    if args.batch:
        stats = converter.batch_convert(args.batch, args.pattern)
        converter.print_enhanced_summary()
        
    elif args.input:
        converter.convert_document(args.input)
        converter.print_enhanced_summary()
        
    else:
        parser.print_help()

if __name__ == "__main__":
    main()
